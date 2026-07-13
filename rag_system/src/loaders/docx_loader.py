import os
from pathlib import Path

from .base import BaseLoader, Document


class DocxLoader(BaseLoader):
    def load(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        doc_result = Document(
            text=text,
            metadata={
                "source_file": path.name,
                "source_type": "docx",
                "document_id": str(path.stem),
                "created_at": self._get_created_at(path),
                "title": path.stem,
                "tags": [],
            },
        )
        return [doc_result]

    @staticmethod
    def _get_created_at(path: Path) -> str:
        try:
            import datetime
            mtime = os.path.getmtime(path)
            return datetime.datetime.fromtimestamp(mtime).isoformat()
        except Exception:
            return ""
